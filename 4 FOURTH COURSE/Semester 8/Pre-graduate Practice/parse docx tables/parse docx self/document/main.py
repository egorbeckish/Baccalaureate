import docx

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from docx.oxml.text.paragraph import CT_P as omxl_paragraph
from docx.oxml.table import CT_Tbl as omxl_table
from docx.oxml.section import CT_SectPr as omxl_section

import regex


def open_docx(path):
	return Document(path)


document = open_docx('test1_tables.docx')
# print(document)


def get_paragraphs(document):
	return document.paragraphs


paragraphs = get_paragraphs(document)
# print(paragraphs)


def get_text(paragraph):
	return paragraph.text


def iter_paragraphs(paragraphs):
	for paragraph in paragraphs:
		print(get_text(paragraph))


# iter_paragraphs(paragraphs)


def get_tables(document):
	return document.tables


tables = get_tables(document)
# print(tables)


def iter_table(table):
	for row in table.rows:
		_row = []
		for cell in row.cells:
			_row += [' '.join(cell.text.split())]

		print('│'.join(_row))


# взяли вторую
# таблицу из списка
table = tables[4]
# iter_table(table)


def get_table(tables):
	table = tables[4]
	_table = []
	for row in table.rows:
		_row = []
		for cell in row.cells:
			_row += [' '.join(cell.text.split())]

		_table += [_row]

	return _table


def get_body(document):
	body = []
	for element in document.element.body:
		if isinstance(element, omxl_paragraph) or isinstance(element, omxl_table):
			body += [element]

	return body


body = get_body(document)
# print(body)

table = get_table(tables)
rows = len(table)
columns = len(table[0])

def lenght_every_column(table, rows, columns):
	dict_lenght_columns = {}

	for i in range(columns):
		maxim_lenght = 0
		for j in range(rows):
			maxim_lenght = max(maxim_lenght, len(table[j][i]))

		dict_lenght_columns[
		    f'столбец_{i}'] = maxim_lenght + 9 if maxim_lenght % 2 else maxim_lenght + 8

	return dict_lenght_columns


# print(lenght_every_column(table, rows, columns))


def show_table(table):
	for row in table:
		# print('│'.join(row))
		print(row)


def format_columns(table, rows, columns):
	dict_lenght_columns = lenght_every_column(table, rows, columns)

	for i in range(columns):
		format_space = dict_lenght_columns[f'столбец_{i}']
		for j in range(rows):
			table[j][i] = f'{table[j][i]:^{format_space}}'


def join_element_rows(table, rows):
	for i in range(rows):
		table[i] = f"│{'│'.join(table[i])}│"
		
		
format_columns(table, rows, columns)
join_element_rows(table, rows)
# show_table(table)
# print(table[0])


def sep_rows(index_sep):
    return f"\n├{'┼'.join(list(map(lambda x: '─' * (x - 1), index_sep)))}┤\n"


def sep_begin_row(index_sep: list[int]) -> str:
    return f"├{'┬'.join(list(map(lambda x: '─' * (x - 1), index_sep)))}┤"

def sep_last_row(index_sep: list[int]) -> str:
    return f"└{'┴'.join(list(map(lambda x: '─' * (x - 1), index_sep)))}┘"


def index_sep_rows(row, columns):
	index_sep = []
	count = 0

	while count != columns:
		index = row[1:].index('│')
		index_sep += [index + 1]
		row = row[index + 1:]
		count += 1

	return index_sep

index_sep = index_sep_rows(table[0], columns)
# print(index_sep)
# print(f"{sep_begin_row(index_sep)}\n{f'{sep_rows(index_sep)}'.join(table)}\n{sep_last_row(index_sep)}")

for element in body:
    if isinstance(element, omxl_paragraph):
        text = Paragraph(element, document).text
        if text:
            print(text)
            print()
    else:
        table = Table(element, document)
        table = [[' '.join(cell.text.split()) for cell in row.cells] for row in table.rows]
        rows = len(table)
        columns = len(table[0])
        format_columns(table, rows, columns)
        join_element_rows(table, rows)
        index_sep = index_sep_rows(table[0], columns)
        print(f"{sep_begin_row(index_sep)}\n{f'{sep_rows(index_sep)}'.join(table)}\n{sep_last_row(index_sep)}")
        print()
