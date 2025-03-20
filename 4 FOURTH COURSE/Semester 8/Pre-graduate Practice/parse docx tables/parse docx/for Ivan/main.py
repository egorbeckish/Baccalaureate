import docx
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

from tabulate import tabulate

import os

import regex

from pdf2docx import Converter, parse

from pprint import *


class ParseWordTable:

    def __init__(self, title: str, obj_table: docx.table.Table | list[docx.table.Table]) -> None:
        self.__obj_table: docx.table.Table = obj_table
        self.__title: str = title
        self.__subtitles: list[list[str]] = self.__slice_obj_table[0]
        self.__table: list[list[str]] = self.__slice_obj_table[1]
        self.__fulltable: list[list[str]] = self.__subtitles + self.__table
        self.__correct_data(self.__fulltable)
        self.__rows: int = self.__table_rows
        self.__columns: int = self.__table_columns


    @property
    def title(self) -> str:
        return self.__title


    @property
    def subtitles(self) -> list[list[str]]:
        return self.__subtitles


    @property
    def table(self) -> list[list[str]]:
        return self.__table


    @property
    def fulltable(self) -> list[list[str]]:
        return self.__fulltable


    @property
    def rows(self) -> int:
        return self.__rows


    @property
    def columns(self) -> int:
        return self.__columns


    def __gauss(self, table: list[list[str]]) -> int:
        for i, row in enumerate(table):
            row_isdigit: list[bool] = map(lambda x: x.split()[0].isdigit() if x else False, row)
            if all(row_isdigit):
                length_row: int = len(row)
                if sum(map(int, row)) == (length_row * (length_row + 1)) / 2 and length_row == len(set(row)):
                    return i
                    table.pop(i)
                    return


    def __delete_space(self, data: list[str]) -> None:
        while '' in data:
            data.pop(data.index(''))


    def __correct_data(self, table: list[list[str]]) -> None:
        if table[0] == table[-1]:
            table.pop(-1)

        for i, row in enumerate(table):
            row = list(map(lambda x: x.split(' '), row))
            for j in range(len(row)):
                self.__delete_space(row[j])
                row[j]: str = ' '.join(row[j])

            table[i]: list[str] = row


    def __join_previously_row(self, table: list[list[str]], join_row: list[list[str]]) -> None:
        if join_row[0][0] == '':
            for i, el in enumerate(join_row[0]):
                if el:
                    table[-1][i] += join_row[0][i]

            join_row.pop(0)


    @property
    def __slice_obj_table(self) -> tuple[list[list[str]], list[list[str]]]:
        if isinstance(self.__obj_table, list):
            layers_table: list[list[list[str]]] = []

            for i, table in enumerate(self.__obj_table):
                data: list[list[str]] = ParseWordTable(self.__title, table)
                if not i:
                    layers_table += data.fulltable

                else:
                    data: list[list[str]] = data.table
                    self.__join_previously_row(layers_table, data)
                    layers_table += data

            index: int = self.__gauss(layers_table)
            return layers_table[:index], layers_table[index + 1:] if index else layers_table[:1], layers_table[1:]


        table: list[list[str]] = [[cell.text for cell in row.cells] for row in self.__obj_table.rows]
        index: int = self.__gauss(table)
        return table[:index], table[index + 1:] if index else table[:1], table[1:]


    @property
    def __table_rows(self) -> int:
        return len(self.__fulltable) + 1


    @property
    def __table_columns(self) -> int:
        return max(map(lambda x: len(x), self.__table))


    @property
    def to_string(self) -> str:
        tmp_table: str = tabulate(
            self.__fulltable,
            [self.__title],
            tablefmt="simple_grid",
            stralign='center',
            showindex=False,
        )

        right_edge: int = tmp_table.index('┐')
        tmp_title: str = ''.join(['╭', '─' * (right_edge - 1), '╮'])
        title: str = f"{tmp_title}\n{f'│{self.__title:^{right_edge - 1}}│'}"

        left_edge: int = tmp_table.index('├')
        table: str = tmp_table[left_edge:].replace('┼', '┬', self.__columns - 1)

        return f'{title}\n{table}'


    def __getitem__(self, index: int) -> str | list[str]:
        return self.__fulltable[index - 1] if index >= 1 \
                else self.__title if self.__rows == abs(index) or not index \
                else self.__fulltable[index]


class WordTable2String:

    def __init__(self, path: str) -> None:
        self.__flag_pdf: bool = False
        if regex.search('.pdf', path):
            self.__flag_pdf: bool = True
            parse(path, 'pdf.docx')

        self.__document: Document = self.__open_docx(path if not self.__flag_pdf else 'pdf.docx')
        self.__body = self.__new_body(self.__document.element.body)
        self.__tables: dict[str: ParseWordTable] = dict()
        self.__table_to_string


    @property
    def tables(self) -> list[str]:
        return self.__tables


    def __open_docx(self, path: str) -> Document:
        return Document(path)


    def __convert_oxml_text_to_string(self, oxml_text: CT_P) -> str:
        return docx.text.paragraph.Paragraph(oxml_text, self.__document).text.lower()


    def __convert_oxml_text_to_string_title(self, oxml_list: list[CT_P]) -> str:
        return '. '.join(map(lambda x: docx.text.paragraph.Paragraph(x, self.__document).text.lower(), oxml_list))


    def __convert_string_to_oxml(self, title: str) -> CT_P:
        tmp_doc: Document = Document()
        paragraph = tmp_doc.add_paragraph(title)
        return paragraph._element


    def __convert_oxml_table_to_table(self, oxml_table: CT_Tbl) -> docx.table.Table:
        return docx.table.Table(oxml_table, self.__document)


    def __delete_space(self, data: list[str]) -> None:
        while '' in data:
            data.pop(data.index(''))


    def __format_title(self, title: str) -> str:
        format_title: list[str] = regex.findall(r'таблица\s[0-9]{1,}\.', title)
        if format_title:
            return title

        tmp_title: str = regex.findall(r'таблица\s[0-9]{1,}', title)[0]
        tmp_text: str = regex.split(r'\s', regex.split(tmp_title, title)[-1])
        self.__delete_space(tmp_text)
        tmp_text: str = ' '.join(tmp_text)

        return tmp_title + '. ' + tmp_text


    def __gauss(self, table: list[list[str]]) -> bool:
        for i, row in enumerate(table):
            row_isdigit: list[bool] = map(lambda x: x.isdigit(), row)
            if all(row_isdigit):
                length_row: int = len(row)
                if sum(map(int, row)) == (length_row * (length_row + 1)) / 2 and length_row == len(set(row)):
                    return True
        return False


    def __new_body(self, body) -> list[docx.oxml]:

        new_body: list[docx.oxml] = []
        index_title_table: int | None = None
        unic_index_title: list[str] = []

        for index, element in enumerate(body):
            if isinstance(element, CT_P):
                text: str = self.__convert_oxml_text_to_string(element)
                regex_text: list[str] = regex.findall(r'таблица\s[0-9]{1,}', text)
                if regex_text:
                    index_title_table: int = index

            if isinstance(element, CT_Tbl):
                table = self.__convert_oxml_table_to_table(element)
                table: list[list[str]] = [[cell.text[::-1].replace(' ', '', 1)[::-1] for cell in row.cells] for row in table.rows]

                if 'ТИПОВЫЕ ТРЕБОВАНИЯ' in table[0][0] or len(table) <= 3:
                    continue

                if index_title_table:
                    if index_title_table not in unic_index_title:
                        title: str = self.__convert_oxml_text_to_string_title(body[index_title_table:index])
                        title: CT_P = self.__convert_string_to_oxml(title)
                        new_body += [title]

                    unic_index_title += [index_title_table]

                new_body += [element]

        new_body: list[CT_P | CT_Tbl | list[CT_Tbl]] = self.__layers_table(new_body)
        return new_body


    def __layers_table(self, body) -> list[CT_P | CT_Tbl | list[CT_Tbl]]:
        index_paragraphs = [i for i in range(len(body)) if isinstance(body[i], CT_P)]
        if not index_paragraphs:
            return []
          
        new_body: list[CT_P | CT_Tbl | list[CT_Tbl]] = []
        for i in range(len(index_paragraphs) - 1):
            element: list[CT_Tbl] = body[index_paragraphs[i] + 1:index_paragraphs[i + 1]]

            new_body += [body[index_paragraphs[i]]]
            if len(element) == 1:
                new_body += element

            else:
                new_body += [element]

        new_body += [body[index_paragraphs[-1]]]
        element: list[CT_Tbl] = body[index_paragraphs[-1] + 1:]
        if len(element) == 1:
            new_body += element

        else:
            new_body += [element]

        return new_body

    @property
    def __table_to_string(self) -> None:
        index_title_table: int | None = None
        for index, element in enumerate(self.__body):
            if isinstance(element, CT_P):
                text: str = self.__convert_oxml_text_to_string(element)
                # regex_text: list[str] = regex.findall(r'таблица\s[0-9]{1,}', text)
                # if regex_text:
                #     index_title_table: int = index

            if isinstance(element, CT_Tbl) or isinstance(element, list):
                title: str = self.__format_title(text)

                if isinstance(element, list):
                    for i, el in enumerate(element):
                        element[i]: docx.table.Table = self.__convert_oxml_table_to_table(el)
                else:
                    element: docx.table.Table = self.__convert_oxml_table_to_table(element)

                self.__tables[title] = ParseWordTable(title, element)

                # if index_title_table:
                #     title: str = self.__convert_oxml_text_to_string_title(self.__body[index_title_table:index])
                #     element: docx.table.Table = self.__convert_oxml_table_to_table(element)
                #     self.__tables[title] = ParseWordTable(title, element)


    def __getitem__(self, index: int) -> str:
        return list(self.__tables.values())[index - 1]



a = WordTable2String('название_файла.docx')

"""
тут будут показаны все обработанные таблицы, если их нет, значит файл без таблиц
"""

pprint(
    a.tables
)

"""
через index можно обращаться к конкретной таблице.
index соответсвует номеру таблицы, т.е. index = 1, тогда будет выведина таблица 1 и т.д
"""

print(
    a[1].to_string
)
